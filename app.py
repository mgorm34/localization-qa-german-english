"""
German↔English Technical Translation QA Pipeline — Flask version
Domain-specific glossaries, confidence highlighting, document upload
"""

import difflib
import html as html_lib
import io
import json
import math
import os
import re
import subprocess
import sys
import unicodedata

import requests as http_requests
import sacrebleu
import spacy
import torch
from flask import Flask, jsonify, request, Response, send_file
from transformers import (
    AutoModelForCausalLM,
    AutoTokenizer,
    MarianMTModel,
    MarianTokenizer,
)

# ── Model loading ────────────────────────────────────────────────────────────

print("Loading models …")

DE_EN_NAME = "Helsinki-NLP/opus-mt-de-en"
EN_DE_NAME = "Helsinki-NLP/opus-mt-en-de"
GPT2_NAME = "gpt2"

de_en_tok = MarianTokenizer.from_pretrained(DE_EN_NAME)
de_en_model = MarianMTModel.from_pretrained(DE_EN_NAME)

en_de_tok = MarianTokenizer.from_pretrained(EN_DE_NAME)
en_de_model = MarianMTModel.from_pretrained(EN_DE_NAME)

gpt2_tok = AutoTokenizer.from_pretrained(GPT2_NAME)
gpt2_model = AutoModelForCausalLM.from_pretrained(GPT2_NAME)

# spaCy for tokenization / lemmatization (used in glossary matching)
def _load_spacy():
    try:
        import spacy_transformers  # noqa: F401
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "spacy-transformers", "-q"])
    try:
        return spacy.load("en_core_web_trf"), "en_core_web_trf"
    except Exception:
        pass
    try:
        return spacy.load("en_core_web_sm"), "en_core_web_sm"
    except OSError:
        subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
        return spacy.load("en_core_web_sm"), "en_core_web_sm"

nlp_en, SPACY_MODEL = _load_spacy()

# German spaCy for source-side glossary matching
try:
    nlp_de = spacy.load("de_core_news_sm")
except OSError:
    subprocess.check_call([sys.executable, "-m", "spacy", "download", "de_core_news_sm"])
    nlp_de = spacy.load("de_core_news_sm")

print(f"All models loaded ✓  (spaCy EN: {SPACY_MODEL})")


# ══════════════════════════════════════════════════════════════════════════════
# DEEPL API
# ══════════════════════════════════════════════════════════════════════════════

DEEPL_API_KEY = os.environ.get("DEEPL_API_KEY", "").strip()
# Use free API if key ends with :fx, otherwise use pro API
if DEEPL_API_KEY.endswith(":fx"):
    DEEPL_URL = "https://api-free.deepl.com/v2/translate"
else:
    DEEPL_URL = "https://api.deepl.com/v2/translate"

if DEEPL_API_KEY and DEEPL_API_KEY != "YOUR_KEY_HERE":
    print(f"DeepL API key found ({len(DEEPL_API_KEY)} chars, ending: ...{DEEPL_API_KEY[-4:]})")
    print(f"DeepL endpoint: {DEEPL_URL}")
else:
    print("WARNING: No DeepL API key found. Will use MarianMT fallback.")
    print(f"  Checked env var DEEPL_API_KEY = '{os.environ.get('DEEPL_API_KEY', '<not set>')}'")
    print(f"  Set DEEPL_API_KEY environment variable or HuggingFace Space secret.")

DEEPL_LANG_MAP = {
    "de-en": ("DE", "EN"),
    "en-de": ("EN", "DE"),
}


def translate_deepl(text, direction="de-en"):
    """Translate using DeepL API. Returns translated text or None on failure."""
    if not DEEPL_API_KEY or DEEPL_API_KEY == "YOUR_KEY_HERE":
        print("DeepL skipped: no API key")
        return None
    src_lang, tgt_lang = DEEPL_LANG_MAP.get(direction, ("DE", "EN"))
    try:
        resp = http_requests.post(DEEPL_URL,
            headers={"Authorization": f"DeepL-Auth-Key {DEEPL_API_KEY}",
                     "Content-Type": "application/json"},
            json={"text": [text], "source_lang": src_lang, "target_lang": tgt_lang},
            timeout=30)
        if resp.status_code == 200:
            return resp.json()["translations"][0]["text"]
        else:
            print(f"DeepL error {resp.status_code}: {resp.text[:200]}")
            return None
    except Exception as e:
        print(f"DeepL request failed: {e}")
        return None


def translate_primary(text, direction="de-en"):
    """Primary translation: DeepL with MarianMT fallback."""
    result = translate_deepl(text, direction)
    if result:
        return result, "deepl"
    # Fallback to MarianMT
    best, _, _ = translate_with_beams(text, direction, fast=True)
    return best, "marian"


# ══════════════════════════════════════════════════════════════════════════════
# WORD ALIGNMENT
# ══════════════════════════════════════════════════════════════════════════════

def compute_word_alignment(source_text, translation_text, direction="de-en"):
    """
    Build bidirectional word alignment using back-translation as a bridge.
    1. Back-translate the target → get pseudo-source
    2. Align pseudo-source tokens to original source tokens (same language, difflib works well)
    3. Use positional correspondence: source[i] ↔ pseudo_source[k] ↔ target[k]
    """
    src_words = source_text.split()
    tgt_words = translation_text.split()

    if not src_words or not tgt_words:
        return {"src_to_tgt": {}, "tgt_to_src": {}}

    src_to_tgt = {i: [] for i in range(len(src_words))}
    tgt_to_src = {i: [] for i in range(len(tgt_words))}

    # Step 1: Back-translate target to get pseudo-source (same language as source)
    reverse_dir = "en-de" if direction == "de-en" else "de-en"
    back_tok, back_model, _, _ = get_models(reverse_dir)
    back_inputs = back_tok(translation_text, return_tensors="pt", padding=True, truncation=True)
    with torch.no_grad():
        back_out = back_model.generate(**back_inputs, num_beams=4, max_length=512)
    pseudo_source = back_tok.decode(back_out[0], skip_special_tokens=True)
    pseudo_words = pseudo_source.split()

    # Step 2: Align original source ↔ pseudo-source (same language — difflib is reliable)
    src_clean = [w.lower().strip(".,;:!?\"'()[]{}") for w in src_words]
    pseudo_clean = [w.lower().strip(".,;:!?\"'()[]{}") for w in pseudo_words]

    # Map pseudo-source indices to original source indices
    pseudo_to_src = {}
    sm = difflib.SequenceMatcher(None, pseudo_clean, src_clean)
    for op, p1, p2, s1, s2 in sm.get_opcodes():
        if op == "equal":
            for offset in range(p2 - p1):
                pseudo_to_src[p1 + offset] = s1 + offset

    # Step 3: Build target ↔ pseudo-source positional mapping
    # Target and pseudo-source come from the same model pass, so they share rough
    # positional correspondence. Use ratio mapping with a window.
    if len(pseudo_words) > 0:
        ratio_tp = len(pseudo_words) / len(tgt_words)
        for j in range(len(tgt_words)):
            # Estimate which pseudo-source word corresponds to target word j
            p_est = min(int(j * ratio_tp), len(pseudo_words) - 1)
            # Check window around estimate
            best_p = p_est
            best_score = -1
            tgt_clean_j = tgt_words[j].lower().strip(".,;:!?\"'()[]{}")
            for p in range(max(0, p_est - 3), min(len(pseudo_words), p_est + 4)):
                # Prefer pseudo words that map to a source word
                score = 1.0 if p in pseudo_to_src else 0.0
                # Boost if pseudo word is similar to any source word nearby
                if p in pseudo_to_src:
                    si = pseudo_to_src[p]
                    score += difflib.SequenceMatcher(None, src_clean[si], tgt_clean_j).ratio()
                if score > best_score:
                    best_score = score
                    best_p = p

            if best_p in pseudo_to_src:
                src_idx = pseudo_to_src[best_p]
                if src_idx not in src_to_tgt or j not in src_to_tgt[src_idx]:
                    src_to_tgt[src_idx].append(j)
                if j not in tgt_to_src or src_idx not in tgt_to_src[j]:
                    tgt_to_src[j].append(src_idx)

    # Step 4: Direct cognate/number matching (catches proper nouns, numbers, shared words)
    for i, sw in enumerate(src_clean):
        if len(sw) < 2:
            continue
        for j, tw in enumerate([w.lower().strip(".,;:!?\"'()[]{}") for w in tgt_words]):
            if sw == tw and j not in src_to_tgt.get(i, []):
                src_to_tgt[i].append(j)
                tgt_to_src[j].append(i)

    # Deduplicate and sort
    for k in src_to_tgt:
        src_to_tgt[k] = sorted(set(src_to_tgt[k]))
    for k in tgt_to_src:
        tgt_to_src[k] = sorted(set(tgt_to_src[k]))

    return {
        "src_to_tgt": {str(k): v for k, v in src_to_tgt.items()},
        "tgt_to_src": {str(k): v for k, v in tgt_to_src.items()},
    }


# ══════════════════════════════════════════════════════════════════════════════
# GLOSSARY SYSTEM
# ══════════════════════════════════════════════════════════════════════════════

# Glossary entry format:
# {
#   "source": "Anspruch",              # German term (or English if EN→DE)
#   "source_lower": "anspruch",        # auto-generated
#   "preferred": "claim",              # preferred translation
#   "alternatives": ["assertion"],     # acceptable alternatives
#   "avoid": ["demand"],               # translations to flag
#   "context": "Patent claims",        # usage note shown to user
#   "confidence": 0.95                 # source reliability (0-1)
# }

DOMAINS = {
    "medical-clinical": {
        "name": "Medical — Clinical",
        "description": "Clinical medicine, diagnosis, patient care",
    },
    "medical-pharma": {
        "name": "Medical — Pharma & Drug Labels",
        "description": "Pharmaceutical, drug labeling, clinical trials",
    },
    "medical-devices": {
        "name": "Medical — Devices",
        "description": "Medical devices, equipment, regulatory",
    },
    "legal-patent": {
        "name": "Legal — Patent Law",
        "description": "Patent applications, claims, IP law",
    },
    "legal-contract": {
        "name": "Legal — Contract Law",
        "description": "Contracts, agreements, commercial law",
    },
    "legal-regulatory": {
        "name": "Legal — Regulatory & Compliance",
        "description": "Regulatory filings, compliance, EU law",
    },
    "history-academic": {
        "name": "History — Academic",
        "description": "Historiography, academic papers, analysis",
    },
    "history-archival": {
        "name": "History — Archival / Primary Sources",
        "description": "Historical documents, treaties, archival material",
    },
    "engineering": {
        "name": "Engineering / Technical",
        "description": "Mechanical, electrical, manufacturing",
    },
    "finance": {
        "name": "Finance & Insurance",
        "description": "Banking, insurance, audit, financial reporting",
    },
    "general": {
        "name": "General",
        "description": "No domain-specific glossary",
    },
}

# In-memory glossary store: domain_id -> [entry, entry, ...]
glossaries = {d: [] for d in DOMAINS}

# Precomputed lookup indices for fast matching
# domain_id -> { source_lower: entry }
glossary_source_index = {d: {} for d in DOMAINS}
# domain_id -> { preferred_lower: entry, alt_lower: entry, ... }
glossary_target_index = {d: {} for d in DOMAINS}

try:
    _base = os.path.dirname(os.path.abspath(__file__))
except NameError:
    _base = os.getcwd()  # Colab / interactive
GLOSSARY_DIR = os.path.join(_base, "glossaries")
os.makedirs(GLOSSARY_DIR, exist_ok=True)


def _rebuild_index(domain):
    """Rebuild lookup indices for a domain after glossary changes."""
    entries = glossaries.get(domain, [])
    src_idx = {}
    tgt_idx = {}
    for entry in entries:
        src_idx[entry["source_lower"]] = entry
        tgt_idx[entry["preferred"].lower()] = entry
        for alt in entry.get("alternatives", []):
            tgt_idx[alt.lower()] = entry
    glossary_source_index[domain] = src_idx
    glossary_target_index[domain] = tgt_idx


def load_glossary(domain, entries):
    """Load a list of glossary entries for a domain."""
    processed = []
    for e in entries:
        entry = dict(e)
        entry["source_lower"] = entry["source"].lower()
        entry.setdefault("alternatives", [])
        entry.setdefault("avoid", [])
        entry.setdefault("context", "")
        entry.setdefault("confidence", 0.8)
        processed.append(entry)
    glossaries[domain] = processed
    _rebuild_index(domain)


def load_glossary_from_file(domain, filepath):
    """Load glossary from a JSON file."""
    with open(filepath, "r", encoding="utf-8") as f:
        entries = json.load(f)
    load_glossary(domain, entries)
    print(f"  Loaded {len(entries)} terms for {domain}")


def save_glossary_to_file(domain):
    """Save current glossary to JSON file."""
    filepath = os.path.join(GLOSSARY_DIR, f"{domain}.json")
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(glossaries[domain], f, indent=2, ensure_ascii=False)


def lookup_source_term(domain, term):
    """Look up a source-language term in the glossary. Returns entry or None."""
    return glossary_source_index.get(domain, {}).get(term.lower())


def lookup_target_term(domain, term):
    """Look up a target-language term in the glossary. Returns entry or None."""
    return glossary_target_index.get(domain, {}).get(term.lower())


def find_glossary_terms_in_source(domain, source_text, direction="de-en"):
    """Find all glossary terms present in the source text. Returns [(start, end, entry), ...]."""
    if domain == "general" or not glossaries.get(domain):
        return []
    src_idx = glossary_source_index.get(domain, {})
    if not src_idx:
        return []

    # Tokenize source
    source_lower = source_text.lower()
    matches = []

    # Check multi-word terms first (longest match)
    sorted_terms = sorted(src_idx.keys(), key=len, reverse=True)
    used_positions = set()

    for term in sorted_terms:
        start = 0
        while True:
            pos = source_lower.find(term, start)
            if pos == -1:
                break
            end = pos + len(term)
            # Check word boundaries
            if (pos == 0 or not source_lower[pos - 1].isalnum()) and \
               (end == len(source_lower) or not source_lower[end].isalnum()):
                # Check no overlap with existing matches
                positions = set(range(pos, end))
                if not positions & used_positions:
                    matches.append((pos, end, src_idx[term]))
                    used_positions |= positions
            start = pos + 1

    return sorted(matches, key=lambda x: x[0])


# Load any existing glossary files on startup
for domain_id in DOMAINS:
    fpath = os.path.join(GLOSSARY_DIR, f"{domain_id}.json")
    if os.path.exists(fpath):
        try:
            load_glossary_from_file(domain_id, fpath)
        except Exception as e:
            print(f"  Warning: failed to load {fpath}: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# TRANSLATION & CONFIDENCE
# ══════════════════════════════════════════════════════════════════════════════

def get_models(direction):
    if direction == "de-en":
        return de_en_tok, de_en_model, en_de_tok, en_de_model
    else:
        return en_de_tok, en_de_model, de_en_tok, de_en_model


def translate_with_beams(text, direction="de-en", fast=False):
    """Translate and return (best, candidates, score_map).
    fast=True: single best translation, no candidates (for document mode).
    """
    fwd_tok, fwd_model, _, _ = get_models(direction)
    inputs = fwd_tok(text, return_tensors="pt", padding=True, truncation=True)

    if fast:
        # Fast mode: minimal beams, 1 result — pure translation, no alternatives
        with torch.no_grad():
            outputs = fwd_model.generate(
                **inputs,
                num_beams=4,
                num_return_sequences=1,
                early_stopping=True,
                max_length=512,
            )
        best = fwd_tok.decode(outputs[0], skip_special_tokens=True)
        return best, [best], {best: 0}

    # Normal mode: 2 passes, 8 beams, 5 sequences (was 4 passes × 30 beams × 20 seqs)
    candidates = []
    seen = set()
    score_map = {}
    global_rank = 0
    for lp in [1.0, 1.2]:
        try:
            with torch.no_grad():
                outputs = fwd_model.generate(
                    **inputs,
                    num_beams=8,
                    num_return_sequences=5,
                    early_stopping=True,
                    max_length=512,
                    length_penalty=lp,
                )
        except Exception:
            continue
        for ids in outputs:
            t = fwd_tok.decode(ids, skip_special_tokens=True)
            if t not in seen:
                seen.add(t)
                candidates.append(t)
                score_map[t] = global_rank
                global_rank += 1
    best = candidates[0] if candidates else ""
    return best, candidates, score_map


def compute_beam_agreement(best, candidates, position):
    """
    How many candidates agree with the best translation at this word position.
    Returns ratio 0.0 - 1.0 (1.0 = all candidates have the same word here).
    """
    best_words = best.split()
    if position >= len(best_words):
        return 1.0
    target = best_words[position].lower()
    agree = 0
    total = 0
    for cand in candidates:
        cand_words = cand.split()
        sm = difflib.SequenceMatcher(None, best_words, cand_words)
        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op == "equal" and i1 <= position < i2:
                agree += 1
                total += 1
                break
            elif op == "replace" and i1 <= position < i2:
                total += 1
                break
        else:
            total += 1  # couldn't align — count as disagreement
    return agree / max(total, 1)


def compute_word_confidence(word, position, best, candidates, domain, direction, source_text):
    """
    Compute confidence for a single word in the translation.
    Combines: beam agreement, glossary match, and word significance.
    Returns (confidence_score 0-100, flags dict).
    """
    flags = {}

    # 1. Beam agreement (0-1)
    beam_agree = compute_beam_agreement(best, candidates, position)

    # 2. Glossary check
    glossary_bonus = 0.0
    entry = lookup_target_term(domain, word)
    if entry:
        if word.lower() == entry["preferred"].lower():
            glossary_bonus = 0.2  # using the preferred term
            flags["glossary"] = "preferred"
            flags["context"] = entry.get("context", "")
        elif word.lower() in [a.lower() for a in entry.get("alternatives", [])]:
            glossary_bonus = 0.05  # acceptable but not preferred
            flags["glossary"] = "acceptable"
            flags["glossary_preferred"] = entry["preferred"]
            flags["context"] = entry.get("context", "")
        else:
            glossary_bonus = 0.0

    # Check if this word is in the avoid list of any glossary entry
    for ent in glossaries.get(domain, []):
        if word.lower() in [a.lower() for a in ent.get("avoid", [])]:
            glossary_bonus = -0.3
            flags["glossary"] = "avoid"
            flags["glossary_preferred"] = ent["preferred"]
            flags["context"] = ent.get("context", "")
            break

    # 3. Word significance — skip function words for highlighting
    doc = nlp_en(word)
    is_content_word = doc[0].pos_ in ("NOUN", "VERB", "ADJ", "ADV", "PROPN") if len(doc) > 0 else False
    flags["content_word"] = is_content_word

    # Combine scores
    raw = beam_agree + glossary_bonus
    # Scale to 0-100
    confidence = max(0, min(100, int(raw * 100)))

    return confidence, flags


def extract_alternatives_filtered(best, candidates, position, score_map, domain):
    """
    Extract alternatives for a word position, filtered to only significant words.
    Returns list of (alt_text, quality, glossary_info) dicts.
    """
    best_words = best.split()
    if position < 0 or position >= len(best_words):
        return []
    target_word = best_words[position].lower()

    # Skip function words — don't offer alternatives for "the", "a", "is", etc.
    doc = nlp_en(best_words[position])
    if len(doc) > 0 and doc[0].pos_ in ("DET", "ADP", "CCONJ", "SCONJ", "PUNCT", "PART", "AUX", "SPACE"):
        return []

    alts = {}
    for cand in candidates[1:]:
        cand_words = cand.split()
        sm = difflib.SequenceMatcher(None, best_words, cand_words)
        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op == "replace" and i1 <= position < i2:
                replacement = " ".join(cand_words[j1:j2])
                if replacement.lower() != target_word:
                    rank = score_map.get(cand, 999)
                    if replacement not in alts or rank < alts[replacement]:
                        alts[replacement] = rank

    total_cands = max(len(candidates), 1)
    result = []
    for alt_text, rank in sorted(alts.items(), key=lambda x: x[1])[:10]:
        quality = max(5, int(100 * (1 - rank / total_cands)))
        # Check glossary for this alternative
        g_info = None
        entry = lookup_target_term(domain, alt_text)
        if entry:
            if alt_text.lower() == entry["preferred"].lower():
                g_info = {"status": "preferred", "context": entry.get("context", "")}
            elif alt_text.lower() in [a.lower() for a in entry.get("alternatives", [])]:
                g_info = {"status": "acceptable", "context": entry.get("context", "")}
        # Check avoid
        for ent in glossaries.get(domain, []):
            if alt_text.lower() in [a.lower() for a in ent.get("avoid", [])]:
                g_info = {"status": "avoid", "preferred": ent["preferred"], "context": ent.get("context", "")}
                break
        result.append({"text": alt_text, "quality": quality, "glossary": g_info})

    return result


def precompute_all_alternatives(best, candidates, score_map, domain):
    """Pre-compute alternatives for every word position, filtered for significance."""
    words = best.split()
    result = {}
    for i in range(len(words)):
        alts = extract_alternatives_filtered(best, candidates, i, score_map, domain)
        if alts:
            result[str(i)] = alts
    return result


def compute_all_confidences(best, candidates, domain, direction, source_text):
    """Compute confidence for every word in the translation."""
    words = best.split()
    confidences = {}
    for i, word in enumerate(words):
        conf, flags = compute_word_confidence(word, i, best, candidates, domain, direction, source_text)
        confidences[str(i)] = {"score": conf, "flags": flags}
    return confidences


def glossary_post_process(translation, domain, direction="de-en"):
    """Force-replace words in translation with glossary preferred terms.
    If a word is in the 'avoid' list of any glossary entry, replace it with preferred.
    """
    if domain == "general" or not glossaries.get(domain):
        return translation

    words = translation.split()
    result = []
    for word in words:
        clean = word.strip(".,;:!?\"'()[]{}").lower()
        replaced = False
        for entry in glossaries.get(domain, []):
            avoid_lower = [a.lower() for a in entry.get("avoid", [])]
            if clean in avoid_lower:
                # Preserve surrounding punctuation
                leading = ""
                trailing = ""
                stripped = word
                while stripped and not stripped[0].isalnum():
                    leading += stripped[0]
                    stripped = stripped[1:]
                while stripped and not stripped[-1].isalnum():
                    trailing = stripped[-1] + trailing
                    stripped = stripped[:-1]
                result.append(leading + entry["preferred"] + trailing)
                replaced = True
                break
        if not replaced:
            result.append(word)
    return " ".join(result)


def constrain_and_retranslate(source_text, chosen_phrase, current_translation, direction="de-en"):
    fwd_tok, fwd_model, back_tok, back_model = get_models(direction)
    inputs = fwd_tok(source_text, return_tensors="pt", padding=True, truncation=True)
    all_candidates = []
    seen = set()
    for lp in [0.6, 0.8, 1.0, 1.2, 1.5]:
        try:
            with torch.no_grad():
                outputs = fwd_model.generate(
                    **inputs, num_beams=50, num_return_sequences=30,
                    early_stopping=True, max_length=512, length_penalty=lp,
                )
        except Exception:
            continue
        for ids in outputs:
            t = fwd_tok.decode(ids, skip_special_tokens=True)
            if t not in seen:
                seen.add(t)
                all_candidates.append(t)
    phrase_lower = chosen_phrase.lower()
    matching = [c for c in all_candidates if phrase_lower in c.lower()]
    if not matching:
        words = current_translation.split()
        best_score, best_spliced = -1, current_translation
        for i in range(len(words)):
            trial_str = " ".join(words[:i] + [chosen_phrase] + words[i + 1:])
            bt_in = back_tok(trial_str, return_tensors="pt", padding=True, truncation=True)
            with torch.no_grad():
                bt_ids = back_model.generate(**bt_in, max_length=512)
            bt = back_tok.decode(bt_ids[0], skip_special_tokens=True)
            sc = sacrebleu.sentence_chrf(bt, [source_text]).score
            if sc > best_score:
                best_score, best_spliced = sc, trial_str
        return best_spliced
    best_score, best_trans = -1, matching[0]
    for cand in matching[:20]:
        bt_in = back_tok(cand, return_tensors="pt", padding=True, truncation=True)
        with torch.no_grad():
            bt_ids = back_model.generate(**bt_in, max_length=512)
        bt = back_tok.decode(bt_ids[0], skip_special_tokens=True)
        sc = sacrebleu.sentence_chrf(bt, [source_text]).score
        if sc > best_score:
            best_score, best_trans = sc, cand
    return best_trans


def score_perplexity(text):
    enc = gpt2_tok(text, return_tensors="pt", truncation=True, max_length=512)
    with torch.no_grad():
        out = gpt2_model(**enc, labels=enc["input_ids"])
    return round(math.exp(out.loss.item()), 2)


def score_backtranslation(source, translation, direction="de-en"):
    _, _, back_tok, back_model = get_models(direction)
    inputs = back_tok(translation, return_tensors="pt", padding=True, truncation=True)
    with torch.no_grad():
        ids = back_model.generate(**inputs, max_length=512)
    back = back_tok.decode(ids[0], skip_special_tokens=True)
    return round(sacrebleu.sentence_chrf(back, [source]).score, 2)


# ── Document extraction helpers ──────────────────────────────────────────────

def extract_text_from_upload(file_bytes, filename):
    """Extract text from uploaded file. Supports .txt, .pdf, .docx."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"])
            import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    elif ext in (".docx", ".doc"):
        try:
            import docx
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
            import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        return file_bytes.decode("utf-8", errors="replace")


def split_into_sentences(text):
    """Split text into sentences for paragraph-by-paragraph translation."""
    # Split on double newlines first (paragraph boundaries)
    paragraphs = re.split(r'\n\s*\n', text)
    # If no paragraph breaks, split on single newlines
    if len(paragraphs) == 1:
        paragraphs = text.split('\n')
    return [p.strip() for p in paragraphs if p.strip()]


# ══════════════════════════════════════════════════════════════════════════════
# FLASK APP
# ══════════════════════════════════════════════════════════════════════════════

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB upload limit
# Allow UTF-8 in JSON responses (umlauts, special chars)
try:
    # Flask 3.x
    app.json.ensure_ascii = False
    app.json.mimetype = "application/json; charset=utf-8"
except AttributeError:
    # Flask 2.x fallback
    app.config['JSON_AS_ASCII'] = False


# CORS: allow HuggingFace iframe to call API across origins
@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

state = {
    "source": "",
    "translation": "",
    "candidates": [],
    "score_map": {},
    "alternatives": {},
    "direction": "de-en",
    "domain": "general",
}


@app.route("/")
def index():
    return Response(HTML_PAGE, mimetype="text/html; charset=utf-8")


@app.route("/api/status", methods=["GET"])
def api_status():
    """Debug endpoint: check engine status, encoding, etc."""
    return jsonify({
        "deepl_configured": bool(DEEPL_API_KEY and DEEPL_API_KEY != "YOUR_KEY_HERE"),
        "deepl_key_length": len(DEEPL_API_KEY) if DEEPL_API_KEY else 0,
        "deepl_key_suffix": DEEPL_API_KEY[-4:] if DEEPL_API_KEY and len(DEEPL_API_KEY) > 4 else "N/A",
        "deepl_url": DEEPL_URL,
        "spacy_model": SPACY_MODEL,
        "glossary_counts": {d: len(glossaries.get(d, [])) for d in DOMAINS},
        "encoding_test": "Ü Ö Ä ß ü ö ä",
    })


@app.route("/api/test-deepl", methods=["GET"])
def api_test_deepl():
    """Test DeepL API directly and return full diagnostic info."""
    import traceback
    result = {"key_loaded": bool(DEEPL_API_KEY), "key_length": len(DEEPL_API_KEY)}
    if not DEEPL_API_KEY:
        result["error"] = "No API key"
        return jsonify(result)
    try:
        resp = http_requests.post(DEEPL_URL,
            headers={"Authorization": f"DeepL-Auth-Key {DEEPL_API_KEY}",
                     "Content-Type": "application/json"},
            json={"text": ["Hallo Welt"], "source_lang": "DE", "target_lang": "EN"},
            timeout=15)
        result["status_code"] = resp.status_code
        result["response_body"] = resp.text[:500]
        if resp.status_code == 200:
            result["translation"] = resp.json()["translations"][0]["text"]
            result["success"] = True
        else:
            result["success"] = False
    except Exception as e:
        result["success"] = False
        result["error"] = str(e)
        result["traceback"] = traceback.format_exc()
    return jsonify(result)


@app.route("/api/domains", methods=["GET"])
def api_domains():
    """Return available domains and their glossary sizes."""
    result = []
    for domain_id, info in DOMAINS.items():
        result.append({
            "id": domain_id,
            "name": info["name"],
            "description": info["description"],
            "term_count": len(glossaries.get(domain_id, [])),
        })
    return jsonify({"domains": result})


@app.route("/api/translate", methods=["POST"])
def api_translate():
    data = request.json
    source = unicodedata.normalize("NFC", data.get("source", "")).strip()
    direction = data.get("direction", "de-en")
    domain = data.get("domain", "general")
    if not source:
        return jsonify({"error": "No source text"}), 400
    if len(source) > 10000:
        return jsonify({"error": "Text exceeds 10,000 character limit"}), 400

    # Primary translation via DeepL (with MarianMT fallback)
    best, engine = translate_primary(source, direction)
    best = glossary_post_process(best, domain, direction)

    # Generate alternatives from MarianMT beam search (lighter pass)
    _, candidates, score_map = translate_with_beams(source, direction)
    alts = precompute_all_alternatives(best, candidates, score_map, domain)
    confidences = compute_all_confidences(best, candidates, domain, direction, source)

    # Word alignment for hover highlighting
    alignment = compute_word_alignment(source, best, direction)

    # Find glossary terms in source
    source_glossary_matches = find_glossary_terms_in_source(domain, source, direction)
    source_matches_data = [
        {"start": s, "end": e, "term": source[s:e], "preferred": ent["preferred"]}
        for s, e, ent in source_glossary_matches
    ]

    # Build source-side glossary annotations (per word)
    src_words = source.split()
    src_glossary = {}
    for s, e, ent in source_glossary_matches:
        # Map character positions to word indices
        char_pos = 0
        for wi, w in enumerate(src_words):
            if char_pos >= s and char_pos < e:
                src_glossary[str(wi)] = {
                    "term": ent["source"],
                    "preferred": ent["preferred"],
                    "context": ent.get("context", ""),
                    "confidence": ent.get("confidence", 0.8),
                }
            char_pos += len(w) + 1

    # Collect review flags (consecutive low-confidence content words)
    review_items = []
    tgt_words = best.split()
    for i, w in enumerate(tgt_words):
        key = str(i)
        conf = confidences.get(key)
        if conf and conf["flags"].get("content_word") and conf["score"] < 50:
            review_items.append({
                "idx": i,
                "word": w,
                "score": conf["score"],
                "reason": conf["flags"].get("glossary", "low_confidence"),
            })

    acc = score_backtranslation(source, best, direction)

    state.update(source=source, translation=best, candidates=candidates,
                 score_map=score_map, alternatives=alts, direction=direction, domain=domain)

    return jsonify({
        "translation": best,
        "words": tgt_words,
        "source_words": src_words,
        "alternatives": alts,
        "confidences": confidences,
        "alignment": alignment,
        "src_glossary": src_glossary,
        "accuracy": acc,
        "engine": engine,
        "source_glossary_matches": source_matches_data,
        "review_items": review_items,
    })


@app.route("/api/retranslate", methods=["POST"])
def api_retranslate():
    data = request.json
    chosen = data.get("chosen", "").strip()
    if not chosen or not state["source"]:
        return jsonify({"error": "Missing data"}), 400

    new_trans = constrain_and_retranslate(
        state["source"], chosen, state["translation"], state["direction"]
    )
    _, new_cands, new_sm = translate_with_beams(state["source"], state["direction"])
    if new_trans not in new_cands:
        new_cands.insert(0, new_trans)
        new_sm[new_trans] = 0
    new_alts = precompute_all_alternatives(new_trans, new_cands, new_sm, state["domain"])
    confidences = compute_all_confidences(new_trans, new_cands, state["domain"], state["direction"], state["source"])
    ppl = score_perplexity(new_trans)
    acc = score_backtranslation(state["source"], new_trans, state["direction"])

    state.update(translation=new_trans, candidates=new_cands,
                 score_map=new_sm, alternatives=new_alts)

    return jsonify({
        "translation": new_trans,
        "words": new_trans.split(),
        "alternatives": new_alts,
        "confidences": confidences,
        "perplexity": ppl,
        "accuracy": acc,
    })


@app.route("/api/upload-document", methods=["POST"])
def api_upload_document():
    """Upload a document, extract text, translate paragraph by paragraph via DeepL."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    direction = request.form.get("direction", "de-en")
    domain = request.form.get("domain", "general")

    file_bytes = f.read()
    text = extract_text_from_upload(file_bytes, f.filename)
    text = unicodedata.normalize("NFC", text)
    if not text.strip():
        return jsonify({"error": "Could not extract text from file"}), 400

    paragraphs = split_into_sentences(text)
    results = []
    for para in paragraphs:
        if len(para.strip()) < 2:
            results.append({"source": para, "translation": para,
                            "source_words": para.split(), "words": para.split(),
                            "alignment": {"src_to_tgt": {}, "tgt_to_src": {}},
                            "alternatives": {}, "confidences": {}})
            continue
        # Use DeepL for speed + quality
        best, engine = translate_primary(para, direction)
        best = glossary_post_process(best, domain, direction)
        alignment = compute_word_alignment(para, best, direction)
        results.append({
            "source": para,
            "translation": best,
            "source_words": para.split(),
            "words": best.split(),
            "alignment": alignment,
            "alternatives": {},
            "confidences": {},
        })

    return jsonify({"paragraphs": results, "filename": f.filename})


@app.route("/api/glossary/<domain_id>", methods=["GET"])
def api_get_glossary(domain_id):
    """Get all terms in a domain glossary."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    return jsonify({"domain": domain_id, "entries": glossaries.get(domain_id, [])})


@app.route("/api/glossary/<domain_id>/add", methods=["POST"])
def api_add_glossary_term(domain_id):
    """Add a term to a domain glossary."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    data = request.json
    entry = {
        "source": data.get("source", ""),
        "source_lower": data.get("source", "").lower(),
        "preferred": data.get("preferred", ""),
        "alternatives": data.get("alternatives", []),
        "avoid": data.get("avoid", []),
        "context": data.get("context", ""),
        "confidence": data.get("confidence", 0.8),
    }
    if not entry["source"] or not entry["preferred"]:
        return jsonify({"error": "source and preferred are required"}), 400
    glossaries.setdefault(domain_id, []).append(entry)
    _rebuild_index(domain_id)
    save_glossary_to_file(domain_id)
    return jsonify({"ok": True, "total": len(glossaries[domain_id])})


@app.route("/api/glossary/<domain_id>/upload", methods=["POST"])
def api_upload_glossary(domain_id):
    """Upload a JSON glossary file to replace/merge with existing."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    f = request.files["file"]
    try:
        entries = json.load(f)
    except Exception as e:
        return jsonify({"error": f"Invalid JSON: {e}"}), 400
    load_glossary(domain_id, entries)
    save_glossary_to_file(domain_id)
    return jsonify({"ok": True, "total": len(glossaries[domain_id])})


@app.route("/api/glossary/<domain_id>/edit/<int:idx>", methods=["PUT"])
def api_edit_glossary_term(domain_id, idx):
    """Edit an existing glossary term by index."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    entries = glossaries.get(domain_id, [])
    if idx < 0 or idx >= len(entries):
        return jsonify({"error": "Index out of range"}), 404
    data = request.json
    entry = entries[idx]
    if "source" in data and data["source"]:
        entry["source"] = data["source"]
        entry["source_lower"] = data["source"].lower()
    if "preferred" in data and data["preferred"]:
        entry["preferred"] = data["preferred"]
    if "alternatives" in data:
        entry["alternatives"] = data["alternatives"] if isinstance(data["alternatives"], list) else [a.strip() for a in data["alternatives"].split(",") if a.strip()]
    if "avoid" in data:
        entry["avoid"] = data["avoid"] if isinstance(data["avoid"], list) else [a.strip() for a in data["avoid"].split(",") if a.strip()]
    if "context" in data:
        entry["context"] = data["context"]
    if "confidence" in data:
        entry["confidence"] = float(data["confidence"])
    _rebuild_index(domain_id)
    save_glossary_to_file(domain_id)
    return jsonify({"ok": True, "entry": entry})


@app.route("/api/glossary/<domain_id>/delete/<int:idx>", methods=["DELETE"])
def api_delete_glossary_term(domain_id, idx):
    """Delete a glossary term by index."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    entries = glossaries.get(domain_id, [])
    if idx < 0 or idx >= len(entries):
        return jsonify({"error": "Index out of range"}), 404
    removed = entries.pop(idx)
    _rebuild_index(domain_id)
    save_glossary_to_file(domain_id)
    return jsonify({"ok": True, "removed": removed["source"], "total": len(entries)})


@app.route("/api/glossary/<domain_id>/export", methods=["GET"])
def api_export_glossary(domain_id):
    """Download glossary as JSON file."""
    if domain_id not in DOMAINS:
        return jsonify({"error": "Unknown domain"}), 404
    entries = glossaries.get(domain_id, [])
    # Strip internal fields for export
    clean = []
    for e in entries:
        clean.append({
            "source": e["source"],
            "preferred": e["preferred"],
            "alternatives": e.get("alternatives", []),
            "avoid": e.get("avoid", []),
            "context": e.get("context", ""),
            "confidence": e.get("confidence", 0.8),
        })
    buf = io.BytesIO()
    buf.write(json.dumps(clean, indent=2, ensure_ascii=False).encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="application/json", as_attachment=True,
                     download_name=f"{domain_id}.json")


@app.route("/glossary")
def glossary_page():
    return Response(GLOSSARY_PAGE, mimetype="text/html")


# ══════════════════════════════════════════════════════════════════════════════
# HTML / CSS / JS FRONTEND
# ══════════════════════════════════════════════════════════════════════════════

HTML_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Technical Translation QA</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: 'Inter', sans-serif; background: #0f172a; color: #e2e8f0; min-height: 100vh; padding: 24px; }
.container { max-width: 1200px; margin: 0 auto; }
h1 { font-size: 22px; font-weight: 700; color: #f1f5f9; margin-bottom: 2px; }
.subtitle { font-size: 13px; color: #94a3b8; margin-bottom: 20px; }

/* Controls */
.controls { display: flex; gap: 12px; align-items: end; flex-wrap: wrap; margin-bottom: 16px; }
.ctrl-group label { display: block; font-size: 11px; font-weight: 600; color: #64748b; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 4px; }
select, .dir-btn { padding: 7px 14px; border: 1px solid #334155; background: #1e293b; color: #e2e8f0; border-radius: 6px; font-family: inherit; font-size: 13px; cursor: pointer; outline: none; }
select:focus { border-color: #3b82f6; }
.dir-btn { font-weight: 500; transition: all 0.15s; }
.dir-btn.active { background: #3b82f6; color: #fff; border-color: #3b82f6; }
.dir-row { display: flex; gap: 6px; }

/* Legend */
.legend { display: flex; gap: 16px; flex-wrap: wrap; font-size: 11px; color: #64748b; margin-bottom: 16px; padding: 8px 12px; background: #1e293b; border-radius: 6px; border: 1px solid #334155; }
.legend-item { display: flex; align-items: center; gap: 4px; }
.legend-swatch { width: 14px; height: 4px; border-radius: 2px; }

/* Input area (before translation) */
.input-area { margin-bottom: 12px; }
.panel-label { font-size: 11px; font-weight: 600; color: #64748b; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 6px; display: flex; justify-content: space-between; }
.char-count { font-weight: 400; color: #475569; }
textarea { width: 100%; min-height: 120px; background: #1e293b; border: 1px solid #334155; border-radius: 8px; color: #e2e8f0; font-family: 'Inter', sans-serif; font-size: 15px; line-height: 1.6; padding: 12px; resize: vertical; outline: none; }
textarea:focus { border-color: #3b82f6; }

/* Side-by-side panels (after translation) */
.side-by-side { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-bottom: 12px; }
@media (max-width: 700px) { .side-by-side { grid-template-columns: 1fr; } }
.text-panel { background: #1e293b; border: 1px solid #334155; border-radius: 8px; padding: 12px; min-height: 120px; line-height: 1.8; font-size: 15px; position: relative; overflow: visible; }

/* Word spans */
.tw-word { color: #e2e8f0; padding: 2px 1px; border-radius: 3px; transition: all 0.15s; cursor: default; }
.tw-has-alt { cursor: pointer; border-bottom: 2px dotted #60a5fa; }
.tw-has-alt:hover { background: #1e3a5f; }
.tw-src-word { cursor: pointer; }

/* Hover alignment highlighting */
.tw-align-highlight { background: rgba(59, 130, 246, 0.25) !important; border-radius: 3px; box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.4); }

/* Confidence highlighting */
.tw-low-conf { background: rgba(251, 146, 60, 0.15); border-bottom: 2px solid #fb923c; }
.tw-very-low-conf { background: rgba(248, 113, 113, 0.15); border-bottom: 2px solid #f87171; }
.tw-glossary-avoid { background: rgba(248, 113, 113, 0.2); border-bottom: 2px solid #ef4444; text-decoration: line-through; text-decoration-color: #ef4444; }
.tw-glossary-preferred { border-bottom: 2px solid #4ade80; }
.tw-src-glossary { border-bottom: 2px solid #4ade80; }

/* Review badge */
.review-badge { display: inline-flex; align-items: center; gap: 4px; padding: 4px 10px; background: #7c2d12; color: #fb923c; font-size: 12px; font-weight: 600; border-radius: 6px; cursor: pointer; transition: all 0.15s; }
.review-badge:hover { background: #9a3412; }

/* Tooltip */
.hover-tooltip { position: fixed; background: #0f172a; border: 1px solid #475569; border-radius: 6px; padding: 8px 12px; font-size: 12px; color: #e2e8f0; z-index: 10000; max-width: 280px; box-shadow: 0 4px 12px rgba(0,0,0,0.5); pointer-events: none; }
.hover-tooltip .tt-label { color: #64748b; font-size: 10px; text-transform: uppercase; letter-spacing: 0.3px; }
.hover-tooltip .tt-term { font-weight: 600; color: #f1f5f9; }
.hover-tooltip .tt-glossary { color: #4ade80; font-size: 11px; margin-top: 4px; }
.hover-tooltip .tt-context { color: #94a3b8; font-size: 11px; font-style: italic; margin-top: 2px; }

.placeholder { color: #64748b; }

/* Dropdown */
.tw-dropdown { position: absolute; background: #1e293b; border: 1px solid #475569; border-radius: 8px; box-shadow: 0 8px 24px rgba(0,0,0,0.5); padding: 4px 0; z-index: 9999; min-width: 160px; max-width: 300px; }
.tw-dropdown button { display: flex; align-items: center; justify-content: space-between; width: 100%; padding: 8px 12px; border: none; background: none; color: #e2e8f0; font-family: 'Inter', sans-serif; font-size: 14px; cursor: pointer; gap: 8px; text-align: left; }
.tw-dropdown button:hover { background: #334155; }
.tw-dropdown .q-score { font-size: 11px; font-weight: 600; min-width: 24px; text-align: right; }
.tw-dropdown .g-tag { font-size: 10px; font-weight: 600; padding: 1px 6px; border-radius: 4px; margin-left: 4px; }
.g-preferred { background: #166534; color: #4ade80; }
.g-acceptable { background: #854d0e; color: #facc15; }
.g-avoid { background: #7f1d1d; color: #f87171; }

/* Scores bar */
.scores-bar { display: flex; gap: 24px; flex-wrap: wrap; margin-top: 10px; padding: 8px 12px; background: #1e293b; border: 1px solid #334155; border-radius: 8px; font-size: 13px; }
.score-label { color: #94a3b8; }
.score-value { font-weight: 600; margin-left: 4px; }
.score-tag { font-size: 12px; margin-left: 4px; }
.score-hint { color: #64748b; font-size: 11px; margin-left: 4px; }

/* Buttons */
.btn-row { display: flex; gap: 10px; margin-bottom: 16px; flex-wrap: wrap; }
.btn { padding: 9px 24px; border: none; border-radius: 8px; font-family: 'Inter', sans-serif; font-size: 14px; font-weight: 600; cursor: pointer; transition: all 0.15s; }
.btn-primary { background: #3b82f6; color: #fff; }
.btn-primary:hover { background: #2563eb; }
.btn-primary:disabled { background: #334155; color: #64748b; cursor: not-allowed; }
.btn-secondary { background: #334155; color: #e2e8f0; }
.btn-secondary:hover { background: #475569; }

/* Doc results: side-by-side per paragraph */
.doc-results { margin-top: 16px; }
.doc-row { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-bottom: 10px; }
@media (max-width: 700px) { .doc-row { grid-template-columns: 1fr; } }
.doc-panel { background: #1e293b; border: 1px solid #334155; border-radius: 8px; padding: 12px; line-height: 1.8; font-size: 14px; position: relative; overflow: visible; }
.doc-panel-label { font-size: 10px; font-weight: 600; color: #475569; text-transform: uppercase; letter-spacing: 0.3px; margin-bottom: 6px; }

/* Loading */
.loading { display: none; }
.loading.active { display: inline-block; }
@keyframes spin { to { transform: rotate(360deg); } }
.spinner { display: inline-block; width: 14px; height: 14px; border: 2px solid #475569; border-top-color: #3b82f6; border-radius: 50%; animation: spin 0.6s linear infinite; vertical-align: middle; margin-left: 6px; }

/* Engine badge */
.engine-badge { font-size: 10px; font-weight: 600; padding: 2px 8px; border-radius: 4px; text-transform: uppercase; letter-spacing: 0.3px; }
.engine-deepl { background: #0f2b46; color: #0fa9e6; }
.engine-marian { background: #1a1a2e; color: #a78bfa; }
</style>
</head>
<body>
<div class="container">
    <h1>Technical Translation QA</h1>
    <p class="subtitle">Domain-specific German↔English translation with confidence highlighting and glossary support.
        <a href="/glossary" style="color:#60a5fa;text-decoration:none;margin-left:12px;font-weight:500;">Manage Glossaries →</a>
    </p>

    <div class="controls">
        <div class="ctrl-group">
            <label>Direction</label>
            <div class="dir-row">
                <button class="dir-btn active" data-dir="de-en" onclick="setDir('de-en')">DE → EN</button>
                <button class="dir-btn" data-dir="en-de" onclick="setDir('en-de')">EN → DE</button>
            </div>
        </div>
        <div class="ctrl-group">
            <label>Domain</label>
            <select id="domainSelect" onchange="currentDomain=this.value"></select>
        </div>
    </div>

    <div class="legend">
        <div class="legend-item"><span class="legend-swatch" style="background:#3b82f6;"></span> Hover alignment</div>
        <div class="legend-item"><span class="legend-swatch" style="background:#60a5fa;"></span> Has alternatives</div>
        <div class="legend-item"><span class="legend-swatch" style="background:#4ade80;"></span> Glossary match</div>
        <div class="legend-item"><span class="legend-swatch" style="background:#fb923c;"></span> Needs review</div>
        <div class="legend-item"><span class="legend-swatch" style="background:#f87171;"></span> Avoid / flagged</div>
    </div>

    <!-- Input area: shown before translation -->
    <div class="input-area" id="inputArea">
        <div class="panel-label">
            <span>Source Text</span>
            <span class="char-count" id="charCount">0 / 10,000</span>
        </div>
        <textarea id="source" placeholder="Enter text here…" maxlength="10000" oninput="updateCharCount()"></textarea>
    </div>

    <!-- Side-by-side: shown after translation -->
    <div class="side-by-side" id="sideBySide" style="display:none;">
        <div>
            <div class="panel-label">
                <span>Source</span>
                <span class="char-count" id="charCount2"></span>
            </div>
            <div class="text-panel" id="sourcePanel"><span class="placeholder">Source</span></div>
        </div>
        <div>
            <div class="panel-label">
                <span>Translation</span>
                <span id="engineBadge"></span>
                <span id="reviewBadgeArea"></span>
            </div>
            <div class="text-panel" id="transBox"><span class="placeholder">Translation will appear here.</span></div>
        </div>
    </div>

    <div id="scoresBar" class="scores-bar" style="display:none;"></div>

    <div class="btn-row">
        <button class="btn btn-primary" id="translateBtn" onclick="doTranslate()">
            Translate <span class="loading" id="transLoading"><span class="spinner"></span></span>
        </button>
        <button class="btn btn-secondary" id="editBtn" onclick="backToEdit()" style="display:none;">Edit Source</button>
        <label class="btn btn-secondary" style="position:relative;">
            Upload Document
            <input type="file" id="fileInput" accept=".txt,.pdf,.docx,.doc" onchange="handleFileUpload(event)" style="position:absolute;opacity:0;width:100%;height:100%;left:0;top:0;cursor:pointer;">
        </label>
    </div>

    <div class="doc-results" id="docResults" style="display:none;"></div>
</div>

<!-- Tooltip element -->
<div class="hover-tooltip" id="hoverTooltip" style="display:none;"></div>

<script>
let currentDir = 'de-en';
let currentDomain = 'general';
let currentData = null;  // full response data for alignment lookups

// Detect base URL: works in iframe (HuggingFace Spaces) and direct access
const API_BASE = (function() {
    // If in an iframe, the actual app is at a different origin
    // Try to detect HuggingFace Space embedding
    const meta = document.querySelector('meta[name="space-host"]');
    if (meta) return 'https://' + meta.content;
    // Check if we're on huggingface.co (iframe wrapper)
    if (window.location.hostname === 'huggingface.co') {
        // Extract space name from URL: /spaces/user/name -> user-name.hf.space
        const m = window.location.pathname.match(/\/spaces\/([^\/]+)\/([^\/]+)/);
        if (m) return 'https://' + m[1] + '-' + m[2] + '.hf.space';
    }
    return '';  // Same origin (direct access or Colab)
})();
console.log('API base URL:', API_BASE || '(same origin)');

// Load domains
async function loadDomains() {
    const res = await fetch(API_BASE + '/api/domains');
    const data = await res.json();
    const sel = document.getElementById('domainSelect');
    sel.innerHTML = '';
    data.domains.forEach(d => {
        const opt = document.createElement('option');
        opt.value = d.id;
        opt.textContent = d.name + (d.term_count > 0 ? ' (' + d.term_count + ' terms)' : '');
        sel.appendChild(opt);
    });
}
loadDomains();

function setDir(dir) {
    currentDir = dir;
    document.querySelectorAll('.dir-btn').forEach(b => b.classList.toggle('active', b.dataset.dir === dir));
}
function updateCharCount() {
    const ta = document.getElementById('source');
    document.getElementById('charCount').textContent = ta.value.length.toLocaleString() + ' / 10,000';
}

function esc(t) { return String(t||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;'); }
function escAttr(t) { return String(t||'').replace(/&/g,'&amp;').replace(/"/g,'&quot;').replace(/</g,'&lt;').replace(/>/g,'&gt;'); }

function qualityColor(q) {
    if (q >= 75) return '#4ade80';
    if (q >= 50) return '#facc15';
    if (q >= 25) return '#fb923c';
    return '#f87171';
}
function accLabel(acc) {
    if (acc > 70) return ['Good', '#4ade80'];
    if (acc > 50) return ['Fair', '#facc15'];
    if (acc > 30) return ['Poor', '#fb923c'];
    return ['Bad', '#f87171'];
}

function closeDropdowns() { document.querySelectorAll('.tw-dropdown').forEach(el => el.remove()); }

// ── Render source words (interactive, hoverable) ──
function renderSourceWords(words, srcGlossary, container, pairId) {
    let html = '';
    words.forEach((word, i) => {
        let cls = 'tw-word tw-src-word';
        let extra = ' data-side="src" data-idx="' + i + '" data-pair="' + pairId + '"';
        const gInfo = srcGlossary && srcGlossary[String(i)];
        if (gInfo) {
            cls += ' tw-src-glossary';
            extra += ' data-glossary="' + escAttr(JSON.stringify(gInfo)) + '"';
        }
        html += '<span class="' + cls + '"' + extra + '>' + esc(word) + '</span> ';
    });
    container.innerHTML = html;
}

// ── Render target words (interactive, hoverable, clickable for alts) ──
function renderTargetWords(words, alts, confidences, container, pairId) {
    let html = '';
    words.forEach((word, i) => {
        const key = String(i);
        const conf = confidences && confidences[key] ? confidences[key] : null;
        const hasAlts = alts && alts[key] && alts[key].length > 0;
        let cls = 'tw-word';
        let extra = ' data-side="tgt" data-idx="' + i + '" data-pair="' + pairId + '"';

        if (conf) {
            const flags = conf.flags || {};
            if (flags.glossary === 'avoid') cls += ' tw-glossary-avoid';
            else if (flags.glossary === 'preferred') cls += ' tw-glossary-preferred';
            else if (conf.score < 30 && flags.content_word) cls += ' tw-very-low-conf';
            else if (conf.score < 60 && flags.content_word) cls += ' tw-low-conf';
        }
        if (hasAlts) {
            cls += ' tw-has-alt';
            extra += ' data-alts="' + escAttr(JSON.stringify(alts[key])) + '"';
        }
        if (conf) extra += ' data-conf="' + escAttr(JSON.stringify(conf)) + '"';

        html += '<span class="' + cls + '"' + extra + '>' + esc(word) + '</span> ';
    });
    container.innerHTML = html;
}

// ── Show side-by-side after translation ──
function showTranslation(data) {
    currentData = data;
    // Hide input, show side-by-side
    document.getElementById('inputArea').style.display = 'none';
    document.getElementById('sideBySide').style.display = 'grid';
    document.getElementById('editBtn').style.display = '';
    document.getElementById('docResults').style.display = 'none';

    // Store alignment data on a global for hover lookups
    window._alignment = data.alignment || {};
    window._pairId = 'main';

    // Render source panel
    renderSourceWords(data.source_words, data.src_glossary || {}, document.getElementById('sourcePanel'), 'main');

    // Render translation panel
    renderTargetWords(data.words, data.alternatives, data.confidences, document.getElementById('transBox'), 'main');

    // Engine badge
    const badge = document.getElementById('engineBadge');
    if (data.engine === 'deepl') {
        badge.innerHTML = '<span class="engine-badge engine-deepl">DeepL</span>';
    } else {
        badge.innerHTML = '<span class="engine-badge engine-marian">MarianMT</span>';
    }

    // Review badge
    const reviewArea = document.getElementById('reviewBadgeArea');
    if (data.review_items && data.review_items.length > 0) {
        reviewArea.innerHTML = '<span class="review-badge" onclick="jumpToReview()">' + data.review_items.length + ' to review</span>';
    } else {
        reviewArea.innerHTML = '';
    }

    // Scores
    if (data.accuracy !== undefined) {
        const [aLbl, aClr] = accLabel(data.accuracy);
        const sb = document.getElementById('scoresBar');
        sb.style.display = 'flex';
        sb.innerHTML =
            '<div><span class="score-label">Accuracy:</span>' +
            '<span class="score-value" style="color:' + aClr + '">' + data.accuracy + '</span>' +
            '<span class="score-tag" style="color:' + aClr + '"> ' + aLbl + '</span>' +
            '<span class="score-hint">(chrF back-translation)</span></div>';
    }
}

function backToEdit() {
    document.getElementById('inputArea').style.display = '';
    document.getElementById('sideBySide').style.display = 'none';
    document.getElementById('editBtn').style.display = 'none';
    document.getElementById('scoresBar').style.display = 'none';
}

function jumpToReview() {
    if (!currentData || !currentData.review_items || !currentData.review_items.length) return;
    const firstIdx = currentData.review_items[0].idx;
    const el = document.querySelector('#transBox .tw-word[data-idx="' + firstIdx + '"]');
    if (el) { el.scrollIntoView({ behavior: 'smooth', block: 'center' }); el.style.outline = '2px solid #fb923c'; setTimeout(() => el.style.outline = '', 2000); }
}

// ── Translate ──
async function doTranslate() {
    const source = document.getElementById('source').value.trim();
    if (!source) return;
    const btn = document.getElementById('translateBtn');
    const spinner = document.getElementById('transLoading');
    btn.disabled = true;
    spinner.classList.add('active');
    try {
        const res = await fetch(API_BASE + '/api/translate', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({source, direction: currentDir, domain: currentDomain}),
        });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        showTranslation(data);
    } catch(e) { alert('Translation failed: ' + e.message);
    } finally { btn.disabled = false; spinner.classList.remove('active'); }
}

// ── Retranslate (pick alternative) ──
async function pickAlternative(phrase) {
    closeDropdowns();
    document.getElementById('transBox').style.opacity = '0.5';
    try {
        const res = await fetch(API_BASE + '/api/retranslate', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({chosen: phrase}),
        });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        // Re-render target panel only
        renderTargetWords(data.words, data.alternatives, data.confidences, document.getElementById('transBox'), 'main');
    } catch(e) { alert('Retranslation failed: ' + e.message);
    } finally { document.getElementById('transBox').style.opacity = '1'; }
}

// ── Document upload ──
async function handleFileUpload(event) {
    const file = event.target.files[0];
    if (!file) return;
    const formData = new FormData();
    formData.append('file', file);
    formData.append('direction', currentDir);
    formData.append('domain', currentDomain);

    const btn = document.getElementById('translateBtn');
    btn.disabled = true;
    btn.textContent = 'Translating document…';
    try {
        const res = await fetch(API_BASE + '/api/upload-document', { method: 'POST', body: formData });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        renderDocResults(data);
    } catch(e) { alert('Upload failed: ' + e.message);
    } finally {
        btn.disabled = false;
        btn.innerHTML = 'Translate <span class="loading" id="transLoading"><span class="spinner"></span></span>';
        event.target.value = '';
    }
}

function renderDocResults(data) {
    document.getElementById('inputArea').style.display = 'none';
    document.getElementById('sideBySide').style.display = 'none';
    document.getElementById('editBtn').style.display = '';
    const container = document.getElementById('docResults');
    container.style.display = 'block';

    let html = '<h3 style="color:#f1f5f9;font-size:16px;margin-bottom:12px;">Document: ' + esc(data.filename) + '</h3>';
    data.paragraphs.forEach((para, idx) => {
        html += '<div class="doc-row" data-doc-pair="doc' + idx + '">';
        html += '<div><div class="doc-panel-label">Source</div><div class="doc-panel" id="docSrc' + idx + '"></div></div>';
        html += '<div><div class="doc-panel-label">Translation</div><div class="doc-panel" id="docTgt' + idx + '"></div></div>';
        html += '</div>';
    });
    container.innerHTML = html;

    // Store alignment data per paragraph
    window._docAlignments = {};
    data.paragraphs.forEach((para, idx) => {
        const pairId = 'doc' + idx;
        window._docAlignments[pairId] = para.alignment || {};
        const srcEl = document.getElementById('docSrc' + idx);
        const tgtEl = document.getElementById('docTgt' + idx);
        if (srcEl) renderSourceWords(para.source_words || para.source.split(' '), {}, srcEl, pairId);
        if (tgtEl) renderTargetWords(para.words, para.alternatives || {}, para.confidences || {}, tgtEl, pairId);
    });
}

// ── Hover alignment logic ──
function getAlignment(pairId) {
    if (pairId === 'main') return window._alignment || {};
    return (window._docAlignments || {})[pairId] || {};
}

function clearAllHighlights() {
    document.querySelectorAll('.tw-align-highlight').forEach(el => el.classList.remove('tw-align-highlight'));
}

const tooltip = document.getElementById('hoverTooltip');

document.addEventListener('mouseover', function(e) {
    const wordEl = e.target.closest('.tw-word');
    if (!wordEl) { clearAllHighlights(); tooltip.style.display = 'none'; return; }

    const side = wordEl.getAttribute('data-side');
    const idx = wordEl.getAttribute('data-idx');
    const pairId = wordEl.getAttribute('data-pair');
    if (!side || idx === null || !pairId) return;

    clearAllHighlights();
    wordEl.classList.add('tw-align-highlight');

    const alignment = getAlignment(pairId);
    let linkedIndices = [];

    if (side === 'src' && alignment.src_to_tgt) {
        linkedIndices = alignment.src_to_tgt[idx] || [];
        linkedIndices.forEach(j => {
            const el = document.querySelector('[data-pair="' + pairId + '"][data-side="tgt"][data-idx="' + j + '"]');
            if (el) el.classList.add('tw-align-highlight');
        });
    } else if (side === 'tgt' && alignment.tgt_to_src) {
        linkedIndices = alignment.tgt_to_src[idx] || [];
        linkedIndices.forEach(j => {
            const el = document.querySelector('[data-pair="' + pairId + '"][data-side="src"][data-idx="' + j + '"]');
            if (el) el.classList.add('tw-align-highlight');
        });
    }

    // Build tooltip
    let ttHtml = '';
    const gData = wordEl.getAttribute('data-glossary');
    const confData = wordEl.getAttribute('data-conf');

    if (gData) {
        try {
            const g = JSON.parse(gData);
            ttHtml += '<div class="tt-label">Glossary</div>';
            ttHtml += '<div class="tt-term">' + esc(g.term) + ' → ' + esc(g.preferred) + '</div>';
            if (g.context) ttHtml += '<div class="tt-context">' + esc(g.context) + '</div>';
        } catch(e) {}
    }
    if (confData) {
        try {
            const c = JSON.parse(confData);
            if (c.flags) {
                if (c.flags.glossary === 'preferred') ttHtml += '<div class="tt-glossary">Glossary preferred term</div>';
                else if (c.flags.glossary === 'acceptable') ttHtml += '<div class="tt-glossary">Acceptable (preferred: ' + esc(c.flags.glossary_preferred) + ')</div>';
                else if (c.flags.glossary === 'avoid') ttHtml += '<div style="color:#f87171;font-size:11px;margin-top:4px;">Avoid — use: ' + esc(c.flags.glossary_preferred) + '</div>';
                if (c.flags.context) ttHtml += '<div class="tt-context">' + esc(c.flags.context) + '</div>';
            }
            if (c.score !== undefined) ttHtml += '<div style="color:#64748b;font-size:11px;margin-top:4px;">Confidence: ' + c.score + '/100</div>';
        } catch(e) {}
    }

    if (ttHtml) {
        tooltip.innerHTML = ttHtml;
        tooltip.style.display = 'block';
        const rect = wordEl.getBoundingClientRect();
        tooltip.style.left = Math.min(rect.left, window.innerWidth - 300) + 'px';
        tooltip.style.top = (rect.bottom + 8) + 'px';
    } else {
        tooltip.style.display = 'none';
    }
});

document.addEventListener('mouseout', function(e) {
    if (!e.relatedTarget || !e.relatedTarget.closest('.tw-word')) {
        clearAllHighlights();
        tooltip.style.display = 'none';
    }
});

// ── Click handler for alternative dropdowns ──
document.addEventListener('click', function(e) {
    var wordEl = e.target.closest('.tw-has-alt');
    if (wordEl) {
        e.stopPropagation();
        closeDropdowns();
        var alts;
        try { alts = JSON.parse(wordEl.getAttribute('data-alts')); } catch(err) { return; }
        if (!alts || alts.length === 0) return;

        var rect = wordEl.getBoundingClientRect();
        var parent = wordEl.closest('.text-panel') || wordEl.closest('.doc-panel');
        if (!parent) return;
        var parentRect = parent.getBoundingClientRect();

        var dd = document.createElement('div');
        dd.className = 'tw-dropdown';
        dd.style.left = (rect.left - parentRect.left) + 'px';
        dd.style.top = (rect.bottom - parentRect.top + 4) + 'px';

        alts.forEach(function(item) {
            var btn = document.createElement('button');
            var txt = document.createElement('span');
            txt.textContent = item.text;
            var right = document.createElement('span');
            right.style.display = 'flex'; right.style.alignItems = 'center'; right.style.gap = '4px';
            if (item.glossary) {
                var tag = document.createElement('span');
                tag.className = 'g-tag';
                if (item.glossary.status === 'preferred') { tag.className += ' g-preferred'; tag.textContent = 'preferred'; }
                else if (item.glossary.status === 'acceptable') { tag.className += ' g-acceptable'; tag.textContent = 'ok'; }
                else if (item.glossary.status === 'avoid') { tag.className += ' g-avoid'; tag.textContent = 'avoid'; }
                right.appendChild(tag);
            }
            var sc = document.createElement('span');
            sc.className = 'q-score'; sc.textContent = item.quality; sc.style.color = qualityColor(item.quality);
            right.appendChild(sc);
            btn.appendChild(txt); btn.appendChild(right);
            btn.onclick = function(ev) { ev.stopPropagation(); pickAlternative(item.text); };
            dd.appendChild(btn);
        });
        parent.appendChild(dd);
        return;
    }
    if (!e.target.closest('.tw-dropdown')) closeDropdowns();
});

// Ctrl+Enter to translate
document.getElementById('source').addEventListener('keydown', function(e) {
    if (e.key === 'Enter' && (e.ctrlKey || e.metaKey)) { e.preventDefault(); doTranslate(); }
});
</script>
</body>
</html>
"""

# ══════════════════════════════════════════════════════════════════════════════
# GLOSSARY MANAGEMENT PAGE
# ══════════════════════════════════════════════════════════════════════════════

GLOSSARY_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Glossary Manager</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
body {
    font-family: 'Inter', sans-serif;
    background: #0f172a;
    color: #e2e8f0;
    min-height: 100vh;
    padding: 24px;
}
.container { max-width: 1200px; margin: 0 auto; }
h1 { font-size: 22px; font-weight: 700; color: #f1f5f9; margin-bottom: 2px; }
.subtitle { font-size: 13px; color: #94a3b8; margin-bottom: 20px; }
a { color: #60a5fa; text-decoration: none; }
a:hover { text-decoration: underline; }

/* Domain cards */
.domain-grid {
    display: grid; grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
    gap: 10px; margin-bottom: 24px;
}
.domain-card {
    background: #1e293b; border: 1px solid #334155; border-radius: 8px; padding: 14px;
    cursor: pointer; transition: all 0.15s;
}
.domain-card:hover { border-color: #475569; background: #263548; }
.domain-card.active { border-color: #3b82f6; background: #1e3a5f; }
.domain-card .d-name { font-size: 14px; font-weight: 600; color: #f1f5f9; margin-bottom: 2px; }
.domain-card .d-desc { font-size: 11px; color: #64748b; margin-bottom: 6px; }
.domain-card .d-count { font-size: 12px; font-weight: 600; color: #3b82f6; }

/* Toolbar */
.toolbar {
    display: flex; gap: 10px; align-items: center; flex-wrap: wrap;
    margin-bottom: 16px; padding: 12px; background: #1e293b; border: 1px solid #334155; border-radius: 8px;
}
.toolbar input[type="text"] {
    flex: 1; min-width: 200px; padding: 8px 12px; background: #0f172a; border: 1px solid #334155;
    border-radius: 6px; color: #e2e8f0; font-family: inherit; font-size: 13px; outline: none;
}
.toolbar input[type="text"]:focus { border-color: #3b82f6; }
.toolbar input[type="text"]::placeholder { color: #475569; }
.btn {
    padding: 8px 16px; border: none; border-radius: 6px; font-family: 'Inter', sans-serif;
    font-size: 13px; font-weight: 600; cursor: pointer; transition: all 0.15s; white-space: nowrap;
}
.btn-primary { background: #3b82f6; color: #fff; }
.btn-primary:hover { background: #2563eb; }
.btn-secondary { background: #334155; color: #e2e8f0; }
.btn-secondary:hover { background: #475569; }
.btn-danger { background: #7f1d1d; color: #f87171; }
.btn-danger:hover { background: #991b1b; }
.btn-sm { padding: 5px 10px; font-size: 12px; }

/* Table */
.table-wrap {
    background: #1e293b; border: 1px solid #334155; border-radius: 8px; overflow: hidden;
}
table {
    width: 100%; border-collapse: collapse; font-size: 13px;
}
thead th {
    background: #0f172a; color: #94a3b8; font-size: 11px; font-weight: 600;
    text-transform: uppercase; letter-spacing: 0.5px; padding: 10px 12px; text-align: left;
    border-bottom: 1px solid #334155; position: sticky; top: 0; z-index: 1;
}
tbody td {
    padding: 10px 12px; border-bottom: 1px solid #1e293b; vertical-align: top;
}
tbody tr { transition: background 0.1s; }
tbody tr:hover { background: #263548; }
tbody tr:last-child td { border-bottom: none; }
.tag {
    display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 11px;
    font-weight: 600; margin: 1px 2px;
}
.tag-alt { background: #1e3a5f; color: #60a5fa; }
.tag-avoid { background: #7f1d1d; color: #f87171; }
.tag-context { color: #64748b; font-size: 12px; font-style: italic; }
.conf-bar {
    width: 50px; height: 6px; background: #334155; border-radius: 3px; display: inline-block;
    vertical-align: middle; margin-right: 6px;
}
.conf-fill { height: 100%; border-radius: 3px; }
.actions-cell { white-space: nowrap; }

/* Modal */
.modal-overlay {
    display: none; position: fixed; top: 0; left: 0; right: 0; bottom: 0;
    background: rgba(0,0,0,0.6); z-index: 1000; align-items: center; justify-content: center;
}
.modal-overlay.show { display: flex; }
.modal {
    background: #1e293b; border: 1px solid #475569; border-radius: 12px; padding: 24px;
    width: 520px; max-width: 95vw; max-height: 90vh; overflow-y: auto;
    box-shadow: 0 16px 48px rgba(0,0,0,0.5);
}
.modal h2 { font-size: 18px; font-weight: 700; color: #f1f5f9; margin-bottom: 16px; }
.form-group { margin-bottom: 14px; }
.form-group label {
    display: block; font-size: 12px; font-weight: 600; color: #94a3b8;
    text-transform: uppercase; letter-spacing: 0.3px; margin-bottom: 4px;
}
.form-group input, .form-group textarea {
    width: 100%; padding: 8px 12px; background: #0f172a; border: 1px solid #334155;
    border-radius: 6px; color: #e2e8f0; font-family: inherit; font-size: 13px; outline: none;
}
.form-group input:focus, .form-group textarea:focus { border-color: #3b82f6; }
.form-group textarea { min-height: 60px; resize: vertical; }
.form-group .hint { font-size: 11px; color: #475569; margin-top: 2px; }
.form-row { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
.modal-actions { display: flex; gap: 8px; justify-content: flex-end; margin-top: 18px; }

/* Empty state */
.empty-state {
    text-align: center; padding: 40px; color: #475569;
}
.empty-state p { margin-bottom: 12px; }

/* Pagination info */
.table-info {
    display: flex; justify-content: space-between; align-items: center;
    padding: 10px 12px; font-size: 12px; color: #64748b;
}
</style>
</head>
<body>
<div class="container">
    <h1>Glossary Manager</h1>
    <p class="subtitle">
        Browse, edit, and build domain-specific translation glossaries.
        <a href="/" style="margin-left:12px;font-weight:500;">← Back to Translator</a>
    </p>

    <div class="domain-grid" id="domainGrid"></div>

    <div id="glossaryPanel" style="display:none;">
        <div class="toolbar">
            <input type="text" id="searchBox" placeholder="Search terms…" oninput="filterTable()">
            <button class="btn btn-primary" onclick="openAddModal()">+ Add Term</button>
            <button class="btn btn-secondary" onclick="exportGlossary()">Export JSON</button>
            <label class="btn btn-secondary" style="position:relative;">
                Import JSON
                <input type="file" accept=".json" onchange="importGlossary(event)"
                    style="position:absolute;opacity:0;width:100%;height:100%;left:0;top:0;cursor:pointer;">
            </label>
        </div>

        <div class="table-wrap">
            <table>
                <thead>
                    <tr>
                        <th style="width:40px;">#</th>
                        <th>Source (DE)</th>
                        <th>Preferred (EN)</th>
                        <th>Alternatives</th>
                        <th>Avoid</th>
                        <th>Context</th>
                        <th style="width:80px;">Conf.</th>
                        <th style="width:100px;">Actions</th>
                    </tr>
                </thead>
                <tbody id="glossaryBody"></tbody>
            </table>
        </div>
        <div class="table-info" id="tableInfo"></div>
    </div>

    <div class="empty-state" id="emptyPrompt">
        <p style="font-size:15px;">Select a domain above to view and edit its glossary.</p>
    </div>
</div>

<!-- Add/Edit Modal -->
<div class="modal-overlay" id="termModal">
    <div class="modal">
        <h2 id="modalTitle">Add Term</h2>
        <div class="form-row">
            <div class="form-group">
                <label>Source (German)</label>
                <input type="text" id="fSource" placeholder="e.g. Befund">
            </div>
            <div class="form-group">
                <label>Preferred Translation</label>
                <input type="text" id="fPreferred" placeholder="e.g. finding">
            </div>
        </div>
        <div class="form-row">
            <div class="form-group">
                <label>Alternatives</label>
                <input type="text" id="fAlternatives" placeholder="e.g. result, observation">
                <div class="hint">Comma-separated acceptable alternatives</div>
            </div>
            <div class="form-group">
                <label>Avoid</label>
                <input type="text" id="fAvoid" placeholder="e.g. report">
                <div class="hint">Comma-separated translations to flag</div>
            </div>
        </div>
        <div class="form-group">
            <label>Context / Usage Note</label>
            <textarea id="fContext" placeholder="e.g. Clinical diagnosis — a diagnostic finding, not a written report"></textarea>
        </div>
        <div class="form-group" style="max-width:200px;">
            <label>Confidence (0.0 – 1.0)</label>
            <input type="number" id="fConfidence" min="0" max="1" step="0.05" value="0.90">
        </div>
        <div class="modal-actions">
            <button class="btn btn-secondary" onclick="closeModal()">Cancel</button>
            <button class="btn btn-primary" id="modalSaveBtn" onclick="saveTerm()">Add Term</button>
        </div>
    </div>
</div>

<script>
// Detect base URL for API calls (same as main page)
const API_BASE = (function() {
    const meta = document.querySelector('meta[name="space-host"]');
    if (meta) return 'https://' + meta.content;
    if (window.location.hostname === 'huggingface.co') {
        const m = window.location.pathname.match(/\/spaces\/([^\/]+)\/([^\/]+)/);
        if (m) return 'https://' + m[1] + '-' + m[2] + '.hf.space';
    }
    return '';
})();

let activeDomain = null;
let allEntries = [];
let editingIdx = null;   // null = adding new, number = editing existing

// ── Load domains ──
async function loadDomains() {
    const res = await fetch(API_BASE + '/api/domains');
    const data = await res.json();
    const grid = document.getElementById('domainGrid');
    grid.innerHTML = '';
    data.domains.forEach(d => {
        if (d.id === 'general') return; // skip general (empty by design)
        const card = document.createElement('div');
        card.className = 'domain-card' + (activeDomain === d.id ? ' active' : '');
        card.onclick = () => selectDomain(d.id);
        card.innerHTML =
            '<div class="d-name">' + esc(d.name) + '</div>' +
            '<div class="d-desc">' + esc(d.description) + '</div>' +
            '<div class="d-count">' + d.term_count + ' terms</div>';
        grid.appendChild(card);
    });
}
loadDomains();

async function selectDomain(domainId) {
    activeDomain = domainId;
    document.querySelectorAll('.domain-card').forEach(c => c.classList.remove('active'));
    document.querySelectorAll('.domain-card').forEach(c => {
        if (c.querySelector('.d-name').textContent.includes(domainId.replace(/-/g,' '))) c.classList.add('active');
    });
    // Re-highlight by re-loading
    await loadDomains();
    await loadGlossary();
    document.getElementById('glossaryPanel').style.display = 'block';
    document.getElementById('emptyPrompt').style.display = 'none';
}

async function loadGlossary() {
    if (!activeDomain) return;
    const res = await fetch(API_BASE + '/api/glossary/' + activeDomain);
    const data = await res.json();
    allEntries = data.entries || [];
    renderTable();
}

function esc(t) {
    if (!t) return '';
    return String(t).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function confColor(c) {
    if (c >= 0.9) return '#4ade80';
    if (c >= 0.8) return '#facc15';
    if (c >= 0.7) return '#fb923c';
    return '#f87171';
}

function renderTable() {
    const search = document.getElementById('searchBox').value.toLowerCase();
    const body = document.getElementById('glossaryBody');
    let html = '';
    let shown = 0;

    allEntries.forEach((e, idx) => {
        // Filter
        if (search) {
            const haystack = [e.source, e.preferred, (e.alternatives||[]).join(' '), (e.avoid||[]).join(' '), e.context||''].join(' ').toLowerCase();
            if (!haystack.includes(search)) return;
        }
        shown++;
        const alts = (e.alternatives || []).map(a => '<span class="tag tag-alt">' + esc(a) + '</span>').join(' ');
        const avoids = (e.avoid || []).map(a => '<span class="tag tag-avoid">' + esc(a) + '</span>').join(' ');
        const conf = e.confidence || 0.8;
        const pct = Math.round(conf * 100);

        html += '<tr>' +
            '<td style="color:#475569;">' + (idx + 1) + '</td>' +
            '<td style="font-weight:600;color:#f1f5f9;">' + esc(e.source) + '</td>' +
            '<td style="color:#4ade80;">' + esc(e.preferred) + '</td>' +
            '<td>' + (alts || '<span style="color:#334155;">—</span>') + '</td>' +
            '<td>' + (avoids || '<span style="color:#334155;">—</span>') + '</td>' +
            '<td><span class="tag-context">' + esc(e.context) + '</span></td>' +
            '<td>' +
                '<span class="conf-bar"><span class="conf-fill" style="width:' + pct + '%;background:' + confColor(conf) + ';"></span></span>' +
                '<span style="color:' + confColor(conf) + ';font-weight:600;font-size:12px;">' + conf.toFixed(2) + '</span>' +
            '</td>' +
            '<td class="actions-cell">' +
                '<button class="btn btn-secondary btn-sm" onclick="openEditModal(' + idx + ')" title="Edit">Edit</button> ' +
                '<button class="btn btn-danger btn-sm" onclick="deleteTerm(' + idx + ')" title="Delete">Del</button>' +
            '</td>' +
            '</tr>';
    });

    if (shown === 0 && search) {
        html = '<tr><td colspan="8" style="text-align:center;padding:24px;color:#475569;">No terms matching "' + esc(search) + '"</td></tr>';
    } else if (shown === 0) {
        html = '<tr><td colspan="8" style="text-align:center;padding:24px;color:#475569;">No terms in this glossary yet. Click <strong>+ Add Term</strong> to start building.</td></tr>';
    }

    body.innerHTML = html;
    document.getElementById('tableInfo').textContent = shown + ' of ' + allEntries.length + ' terms' + (search ? ' (filtered)' : '');
}

function filterTable() { renderTable(); }

// ── Modal ──
function openAddModal() {
    editingIdx = null;
    document.getElementById('modalTitle').textContent = 'Add Term';
    document.getElementById('modalSaveBtn').textContent = 'Add Term';
    document.getElementById('fSource').value = '';
    document.getElementById('fPreferred').value = '';
    document.getElementById('fAlternatives').value = '';
    document.getElementById('fAvoid').value = '';
    document.getElementById('fContext').value = '';
    document.getElementById('fConfidence').value = '0.90';
    document.getElementById('termModal').classList.add('show');
    document.getElementById('fSource').focus();
}

function openEditModal(idx) {
    editingIdx = idx;
    const e = allEntries[idx];
    document.getElementById('modalTitle').textContent = 'Edit Term';
    document.getElementById('modalSaveBtn').textContent = 'Save Changes';
    document.getElementById('fSource').value = e.source || '';
    document.getElementById('fPreferred').value = e.preferred || '';
    document.getElementById('fAlternatives').value = (e.alternatives || []).join(', ');
    document.getElementById('fAvoid').value = (e.avoid || []).join(', ');
    document.getElementById('fContext').value = e.context || '';
    document.getElementById('fConfidence').value = (e.confidence || 0.8).toFixed(2);
    document.getElementById('termModal').classList.add('show');
    document.getElementById('fSource').focus();
}

function closeModal() {
    document.getElementById('termModal').classList.remove('show');
    editingIdx = null;
}

async function saveTerm() {
    const source = document.getElementById('fSource').value.trim();
    const preferred = document.getElementById('fPreferred').value.trim();
    if (!source || !preferred) { alert('Source and Preferred are required.'); return; }

    const payload = {
        source: source,
        preferred: preferred,
        alternatives: document.getElementById('fAlternatives').value.split(',').map(s => s.trim()).filter(Boolean),
        avoid: document.getElementById('fAvoid').value.split(',').map(s => s.trim()).filter(Boolean),
        context: document.getElementById('fContext').value.trim(),
        confidence: parseFloat(document.getElementById('fConfidence').value) || 0.8,
    };

    let url, method;
    if (editingIdx !== null) {
        url = API_BASE + '/api/glossary/' + activeDomain + '/edit/' + editingIdx;
        method = 'PUT';
    } else {
        url = API_BASE + '/api/glossary/' + activeDomain + '/add';
        method = 'POST';
    }

    try {
        const res = await fetch(url, {
            method: method,
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify(payload),
        });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        closeModal();
        await loadGlossary();
        await loadDomains();  // refresh term counts
    } catch(e) { alert('Failed: ' + e.message); }
}

async function deleteTerm(idx) {
    const e = allEntries[idx];
    if (!confirm('Delete "' + e.source + ' → ' + e.preferred + '"?')) return;
    try {
        const res = await fetch(API_BASE + '/api/glossary/' + activeDomain + '/delete/' + idx, { method: 'DELETE' });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        await loadGlossary();
        await loadDomains();
    } catch(e) { alert('Failed: ' + e.message); }
}

// ── Import / Export ──
function exportGlossary() {
    if (!activeDomain) return;
    window.location.href = API_BASE + '/api/glossary/' + activeDomain + '/export';
}

async function importGlossary(event) {
    const file = event.target.files[0];
    if (!file || !activeDomain) return;
    if (!confirm('This will replace the entire "' + activeDomain + '" glossary with the uploaded file. Continue?')) {
        event.target.value = '';
        return;
    }
    const formData = new FormData();
    formData.append('file', file);
    try {
        const res = await fetch(API_BASE + '/api/glossary/' + activeDomain + '/upload', { method: 'POST', body: formData });
        const data = await res.json();
        if (data.error) { alert(data.error); return; }
        await loadGlossary();
        await loadDomains();
    } catch(e) { alert('Import failed: ' + e.message);
    } finally { event.target.value = ''; }
}

// Close modal on overlay click
document.getElementById('termModal').addEventListener('click', function(e) {
    if (e.target === this) closeModal();
});

// Keyboard shortcuts
document.addEventListener('keydown', function(e) {
    if (e.key === 'Escape') closeModal();
});
</script>
</body>
</html>
"""


# ── Run ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import socket

    def _in_colab():
        try:
            from google.colab import output  # noqa: F401
            return True
        except ImportError:
            return False

    if _in_colab():
        # ── Colab: background thread + proxy ──
        import threading

        def find_free_port():
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('', 0))
                return s.getsockname()[1]

        port = find_free_port()
        threading.Thread(target=lambda: app.run(host="0.0.0.0", port=port)).start()

        from google.colab.output import eval_js
        url = eval_js(f"google.colab.kernel.proxyPort({port})")
        print(f"\n✅ App running! Open this URL:\n{url}\n")
    else:
        # ── HuggingFace Spaces / local: standard launch ──
        port = int(os.environ.get("PORT", 7860))
        print(f"\n✅ Starting on port {port} …\n")
        app.run(host="0.0.0.0", port=port)
